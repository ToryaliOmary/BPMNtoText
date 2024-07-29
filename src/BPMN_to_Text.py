import xml.etree.ElementTree as ET
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def extract_bpmn_details(file_path):
    tree = ET.parse(file_path)
    root = tree.getroot()
    ns = {'bpmn': 'http://www.omg.org/spec/BPMN/20100524/MODEL'}

    # Information extrahieren
    process_name = root.find('.//bpmn:process', ns).get('name', 'Unbekannter Prozess')
    author = root.find('.//bpmn:documentation', ns).text if root.find('.//bpmn:documentation', ns) is not None else 'Unbekannter Autor'
    lanes = [lane.get('name') for lane in root.findall('.//bpmn:lane', ns)]

    # Tasks extrahieren
    tasks = []
    for task in root.findall('.//bpmn:task', ns):
        tasks.append({
            'name': task.get('name', 'Unbenannte Aufgabe'),
            'description': task.find('bpmn:documentation', ns).text if task.find('bpmn:documentation', ns) is not None else 'Keine Beschreibung'
        })
    return process_name, author, lanes, tasks

def create_word_document(process_name, author, lanes, tasks, file_path):
    doc = Document()
    doc.add_heading('BPMN Dokumentation', 0)
    doc.add_heading('I. Information', level=1)
    doc.add_paragraph(f'Name des Prozesses: {process_name}')
    doc.add_paragraph(f'Autor: {author}')
    doc.add_paragraph('Beteiligte:')
    for lane in lanes:
        doc.add_paragraph(lane, style='ListBullet')

    doc.add_heading('II. Aufgaben', level=1)
    for task in tasks:
        doc.add_paragraph(task['name'], style='ListNumber')
        doc.add_paragraph(f"Beschreibung: {task['description']}")
    
    doc.save(file_path)

def create_pdf_document(process_name, author, lanes, tasks, file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 750, "BPMN Dokumentation")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(72, 730, "I. Information")
    c.setFont("Helvetica", 12)
    c.drawString(72, 710, f"Name des Prozesses: {process_name}")
    c.drawString(72, 690, f"Autor: {author}")
    c.drawString(72, 670, "Beteiligte:")
    for i, lane in enumerate(lanes, start=1):
        c.drawString(82, 670 - i*20, lane)
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(72, 620, "II. Aufgaben")
    y = 600
    for task in tasks:
        c.setFont("Helvetica-Bold", 12)
        c.drawString(82, y, task['name'])
        c.setFont("Helvetica", 12)
        c.drawString(82, y-20, f"Beschreibung: {task['description']}")
        y -= 40

    c.save()

# Dateipfad zur BPMN XML-Datei
bpmn_file_path = 'C:\\Users\\huyenp\\Desktop\\Huyen\\Lernen\\UNi\\HTW\\WI\\2. Semester\\2. Informationsmanagement\\Test\\export.bpmn'

# Dateipfade für die Ausgabedateien
word_output_path = bpmn_file_path.replace('.bpmn', '.docx')
pdf_output_path = bpmn_file_path.replace('.bpmn', '.pdf')

# Daten aus BPMN extrahieren
process_name, author, lanes, tasks = extract_bpmn_details(bpmn_file_path)

# Word-Dokument erstellen
create_word_document(process_name, author, lanes, tasks, word_output_path)

# PDF-Dokument erstellen
create_pdf_document(process_name, author, lanes, tasks, pdf_output_path)
print('fertig')